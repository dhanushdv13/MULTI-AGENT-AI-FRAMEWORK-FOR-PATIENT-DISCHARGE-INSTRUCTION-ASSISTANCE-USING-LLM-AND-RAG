"""
PDF to Word Document — FAST VERSION

Uses OCR cache from processor.py when available (no double OCR).
For standalone use, falls back to its own OCR.

Digital PDFs: ~1-2 seconds (native text extraction)
Scanned PDFs (with cache): ~1-2 seconds (reads cached OCR)
Scanned PDFs (no cache): ~10-20 seconds (runs OCR)
"""
import sys
import re
import json
import logging
import time
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING

logger = logging.getLogger("pdf_to_word")


# ══════════════════════════════════════════════════════════════════════
# OCR Cache — reads results saved by processor.py
# ══════════════════════════════════════════════════════════════════════

def _load_ocr_cache(pdf_path: str) -> list | None:
    """Load cached OCR results from processor pipeline."""
    cache_path = Path(pdf_path).parent / f"{Path(pdf_path).stem}_ocr_cache.json"
    if cache_path.exists():
        try:
            with open(cache_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return None
    return None


# ══════════════════════════════════════════════════════════════════════
# Native text extraction (for digital PDF pages)
# ══════════════════════════════════════════════════════════════════════

def _extract_native_page(page) -> dict:
    """Extract structured text from a digital PDF page via PyMuPDF."""
    text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
    blocks = []

    for block in text_dict.get("blocks", []):
        if block["type"] != 0:
            continue
        for line in block.get("lines", []):
            spans = []
            for span in line.get("spans", []):
                txt = span.get("text", "").strip()
                if txt:
                    spans.append({
                        "text": txt,
                        "size": span.get("size", 10),
                        "flags": span.get("flags", 0),
                    })
            if spans:
                blocks.append(spans)

    plain = "\n".join(" ".join(s["text"] for s in line) for line in blocks)
    return {"blocks": blocks, "text": plain, "char_count": len(plain)}


# ══════════════════════════════════════════════════════════════════════
# Convert OCR elements to block format (from cache or fresh OCR)
# ══════════════════════════════════════════════════════════════════════

def _cluster_into_lines(elements, y_tolerance=10):
    if not elements:
        return []
    sorted_elems = sorted(elements, key=lambda e: e["y1"])
    lines = []
    cur = [sorted_elems[0]]
    cur_y = sorted_elems[0]["y1"]
    for elem in sorted_elems[1:]:
        if abs(elem["y1"] - cur_y) <= y_tolerance:
            cur.append(elem)
        else:
            cur.sort(key=lambda e: e["x1"])
            lines.append(cur)
            cur = [elem]
            cur_y = elem["y1"]
    cur.sort(key=lambda e: e["x1"])
    lines.append(cur)
    return lines


def _elements_to_blocks(elements):
    """Convert OCR elements (with x1/y1/x2/y2) to block format for Word builder."""
    lines = _cluster_into_lines(elements)
    blocks = []
    for line in lines:
        spans = []
        for e in line:
            h = e.get("y2", e["y1"] + 10) - e["y1"]
            spans.append({"text": e["text"], "size": min(12, max(8, h * 0.7)), "flags": 0})
        blocks.append(spans)
    return blocks


# ══════════════════════════════════════════════════════════════════════
# Fallback OCR (only used when no cache exists)
# ══════════════════════════════════════════════════════════════════════

def _ocr_page_standalone(page_num: int, pdf_path: str, dpi: int = 150) -> dict:
    """OCR a single page (standalone, creates its own OCR engine)."""
    import numpy as np
    from rapidocr_onnxruntime import RapidOCR

    pdf_doc = fitz.open(pdf_path)
    page = pdf_doc[page_num]
    pw, ph = page.rect.width, page.rect.height

    pix = page.get_pixmap(dpi=dpi)
    iw, ih = pix.width, pix.height
    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(ih, iw, pix.n)
    if pix.n == 4:
        img = img[:, :, :3]
    pdf_doc.close()

    ocr = RapidOCR()
    result, _ = ocr(img)
    if not result:
        return {"blocks": [], "text": ""}

    elements = []
    sx, sy = pw / iw, ph / ih
    for det in result:
        bbox, text, conf = det[0], det[1], det[2]
        if conf < 0.3:
            continue
        y1 = min(p[1] for p in bbox) * sy
        x1 = min(p[0] for p in bbox) * sx
        y2 = max(p[1] for p in bbox) * sy
        x2 = max(p[0] for p in bbox) * sx
        elements.append({"text": text, "x1": x1, "y1": y1, "x2": x2, "y2": y2})

    blocks = _elements_to_blocks(elements)
    plain = "\n".join(" ".join(s["text"] for s in line) for line in blocks)
    return {"blocks": blocks, "text": plain}


# ══════════════════════════════════════════════════════════════════════
# Word Document Builder — COMPACT layout
# ══════════════════════════════════════════════════════════════════════

def _is_heading(spans):
    if not spans:
        return False
    text = " ".join(s["text"] for s in spans)
    avg_size = sum(s["size"] for s in spans) / len(spans)
    is_bold = any(s["flags"] & (1 << 4) for s in spans)
    if is_bold and avg_size > 11:
        return True
    if text.isupper() and 3 < len(text.strip()) < 60:
        return True
    return False


def _build_word_doc(all_pages_data, output_path):
    """Build a compact Word document matching original page count."""
    doc = Document()

    # Set compact default style
    style = doc.styles['Normal']
    style.font.size = Pt(9)
    style.font.name = 'Calibri'
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for page_idx, page_data in enumerate(all_pages_data):
        if not page_data:
            continue

        blocks = page_data.get("blocks", [])
        if not blocks:
            # Fallback: use plain text
            text = page_data.get("text", "")
            if text:
                if page_idx > 0:
                    doc.add_page_break()
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        doc.add_paragraph(line)
            continue

        if page_idx > 0:
            doc.add_page_break()

        # Detect table-like rows (multiple spans with significant gaps)
        i = 0
        while i < len(blocks):
            spans = blocks[i]

            # Check for table rows: 3+ spans = likely table
            is_table_row = len(spans) >= 3
            if is_table_row and i + 1 < len(blocks) and len(blocks[i + 1]) >= 3:
                # Collect consecutive table rows
                table_rows = []
                while i < len(blocks) and len(blocks[i]) >= 3:
                    table_rows.append(blocks[i])
                    i += 1

                if len(table_rows) >= 2:
                    max_cols = max(len(r) for r in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=max_cols)
                    table.style = 'Table Grid'
                    for ri, row in enumerate(table_rows):
                        for ci, span in enumerate(row):
                            if ci < max_cols:
                                cell = table.cell(ri, ci)
                                cell.text = span["text"]
                                for p in cell.paragraphs:
                                    p.paragraph_format.space_after = Pt(0)
                                    for run in p.runs:
                                        run.font.size = Pt(8)
                    continue

            # Heading
            if _is_heading(spans):
                text = " ".join(s["text"] for s in spans)
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.bold = True
                run.font.size = Pt(11)
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(2)
            # Key-value pair (e.g., "Name: John Doe")
            elif len(spans) == 2 and spans[0]["text"].rstrip().endswith(":"):
                para = doc.add_paragraph()
                run = para.add_run(spans[0]["text"] + " ")
                run.bold = True
                run.font.size = Pt(9)
                run2 = para.add_run(spans[1]["text"])
                run2.font.size = Pt(9)
            # Normal text line
            else:
                text = " ".join(s["text"] for s in spans)
                if text.strip():
                    para = doc.add_paragraph(text)
                    para.paragraph_format.space_after = Pt(1)

            i += 1

    doc.save(str(output_path))


# ══════════════════════════════════════════════════════════════════════
# Main Function
# ══════════════════════════════════════════════════════════════════════

def pdf_to_word(input_pdf, output_docx=None, dpi=150, max_workers=8):
    """
    Convert PDF to Word — FAST.
    
    Uses OCR cache when available (from processor.py pipeline).
    Falls back to native extraction or fresh OCR.
    """
    start = time.time()

    input_path = Path(input_pdf)
    if output_docx is None:
        output_docx = input_path.parent / f"{input_path.stem}.docx"
    output_path = Path(output_docx)

    logger.info("PDF->Word: %s", input_path.name)

    # Open PDF
    pdf_doc = fitz.open(str(input_path))
    total_pages = len(pdf_doc)

    all_pages_data = [None] * total_pages

    # ── Step 1: Check for OCR cache (from processor pipeline) ──────
    cache = _load_ocr_cache(str(input_path))

    if cache:
        logger.info("  Using OCR cache (no re-OCR!)")
        # For native pages: use PyMuPDF structured extraction (has font info)
        # For OCR pages: use cached elements -> blocks
        for idx, entry in enumerate(cache):
            if idx >= total_pages:
                break
            if entry is None:
                continue

            if entry.get("is_native", False) or not entry.get("elements"):
                # Native page — extract structured text from PDF
                page = pdf_doc[idx]
                native = _extract_native_page(page)
                all_pages_data[idx] = native
            else:
                # Scanned page — convert cached OCR elements to blocks
                blocks = _elements_to_blocks(entry["elements"])
                all_pages_data[idx] = {
                    "blocks": blocks,
                    "text": entry.get("text", ""),
                }

        pdf_doc.close()

    else:
        # ── Step 2: No cache — do our own extraction ──────────────
        logger.info("  No OCR cache found, extracting from scratch...")
        pages_needing_ocr = []

        for pn in range(total_pages):
            page = pdf_doc[pn]
            native = _extract_native_page(page)
            if native["char_count"] >= 30:
                all_pages_data[pn] = native
            else:
                pages_needing_ocr.append(pn)

        pdf_doc.close()

        native_count = total_pages - len(pages_needing_ocr)
        logger.info("  Pages: %d total | %d native | %d need OCR",
                     total_pages, native_count, len(pages_needing_ocr))

        if pages_needing_ocr:
            from concurrent.futures import ProcessPoolExecutor, as_completed as _ac
            workers = min(max_workers, len(pages_needing_ocr))
            logger.info("  Running OCR on %d pages with %d workers...", len(pages_needing_ocr), workers)

            completed = 0
            with ProcessPoolExecutor(max_workers=workers) as executor:
                futures = {
                    executor.submit(_ocr_page_standalone, pn, str(input_path), dpi): pn
                    for pn in pages_needing_ocr
                }
                for future in _ac(futures):
                    pn = futures[future]
                    try:
                        all_pages_data[pn] = future.result()
                    except Exception as e:
                        logger.error("  OCR page %d failed: %s", pn + 1, e)
                        all_pages_data[pn] = {"blocks": [], "text": ""}
                    completed += 1
                    logger.info("  OCR: %d/%d pages done", completed, len(pages_needing_ocr))

    # ── Step 3: Build Word document ────────────────────────────────
    _build_word_doc(all_pages_data, output_path)

    elapsed = time.time() - start
    logger.info("  Word doc saved: %s (%.1fs)", output_path.name, elapsed)

    return str(output_path)


# ══════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s | %(name)-12s | %(message)s",
        datefmt="%H:%M:%S",
    )

    if len(sys.argv) < 2:
        print("Usage: python pdf_to_word.py <input.pdf> [output.docx]")
        sys.exit(1)

    pdf_in = sys.argv[1]
    docx_out = sys.argv[2] if len(sys.argv) >= 3 else None
    output = pdf_to_word(pdf_in, docx_out, dpi=150, max_workers=8)
    print(f"\nWord document created: {output}")
