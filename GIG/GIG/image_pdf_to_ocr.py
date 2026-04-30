"""
PDF to Word Document Reconstruction
Extracts text, tables, and layout from image-based PDFs and creates a Word document.

Usage:
    .venv\Scripts\python.exe pdf_to_word.py tes.pdf output.docx
"""
import os
import sys
import fitz  # PyMuPDF
from rapidocr_onnxruntime import RapidOCR
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import re


def cluster_elements_into_lines(elements, y_tolerance=10):
    """
    Group OCR elements into lines based on Y-coordinate proximity.
    Returns list of lines, each line is list of elements sorted by X.
    """
    if not elements:
        return []
    
    # Sort by Y first
    sorted_elems = sorted(elements, key=lambda e: e['y1'])
    
    lines = []
    current_line = [sorted_elems[0]]
    current_y = sorted_elems[0]['y1']
    
    for elem in sorted_elems[1:]:
        if abs(elem['y1'] - current_y) <= y_tolerance:
            current_line.append(elem)
        else:
            # Sort current line by X and save
            current_line.sort(key=lambda e: e['x1'])
            lines.append(current_line)
            current_line = [elem]
            current_y = elem['y1']
    
    # Don't forget last line
    current_line.sort(key=lambda e: e['x1'])
    lines.append(current_line)
    
    return lines


def detect_table_rows(lines, x_gap_threshold=50, min_columns=2):
    """
    Detect lines that might be table rows (multiple columns with similar X positions).
    Returns list of (line_index, is_table_row, num_columns).
    """
    table_info = []
    
    for i, line in enumerate(lines):
        if len(line) >= min_columns:
            # Check if elements have significant X gaps (columns)
            gaps = []
            for j in range(1, len(line)):
                gap = line[j]['x1'] - line[j-1]['x2']
                gaps.append(gap)
            
            # If there are significant gaps, this might be a table row
            significant_gaps = sum(1 for g in gaps if g > x_gap_threshold)
            if significant_gaps >= 1:
                table_info.append((i, True, len(line)))
            else:
                table_info.append((i, False, 1))
        else:
            table_info.append((i, False, 1))
    
    return table_info


def extract_page_elements(ocr_result, page_width, page_height, img_width, img_height):
    """
    Extract elements from OCR result with normalized coordinates.
    """
    elements = []
    
    if not ocr_result:
        return elements
    
    scale_x = page_width / img_width
    scale_y = page_height / img_height
    
    for detection in ocr_result:
        bbox = detection[0]  # [[x1,y1], [x2,y2], [x3,y3], [x4,y4]]
        text = detection[1]
        confidence = detection[2]
        
        if confidence < 0.3:
            continue
        
        # Get bounding box
        x1 = min([pt[0] for pt in bbox]) * scale_x
        y1 = min([pt[1] for pt in bbox]) * scale_y
        x2 = max([pt[0] for pt in bbox]) * scale_x
        y2 = max([pt[1] for pt in bbox]) * scale_y
        
        # Calculate text height (for font size estimation)
        height = y2 - y1
        
        elements.append({
            'text': text,
            'x1': x1, 'y1': y1, 'x2': x2, 'y2': y2,
            'height': height,
            'confidence': confidence
        })
    
    return elements


def is_heading(text, height, avg_height):
    """
    Heuristic to detect if text is a heading.
    """
    # Larger text
    if height > avg_height * 1.3:
        return True
    # All caps and short
    if text.isupper() and len(text) < 50:
        return True
    # Common heading patterns
    heading_patterns = [
        r'^[A-Z][A-Z\s]+:$',  # ALL CAPS WITH COLON
        r'^[A-Z][A-Z\s]+$',   # ALL CAPS
        r'^\d+\.\s+[A-Z]',    # Numbered heading
    ]
    for pattern in heading_patterns:
        if re.match(pattern, text):
            return True
    return False


def pdf_to_word(input_pdf, output_docx=None, dpi=300):
    """
    Convert PDF to Word document with layout preservation.
    
    Args:
        input_pdf: Path to input PDF
        output_docx: Path to output Word document (default: {input}.docx)
        dpi: Resolution for OCR
    
    Returns:
        output_path
    """
    input_path = Path(input_pdf)
    if output_docx is None:
        output_docx = input_path.parent / f"{input_path.stem}.docx"
    output_path = Path(output_docx)
    
    print("=" * 60)
    print("PDF TO WORD DOCUMENT RECONSTRUCTION")
    print("=" * 60)
    print(f"Input:  {input_path.name}")
    print(f"Output: {output_path.name}")
    print()
    
    # Initialize OCR
    print("Initializing OCR...")
    ocr = RapidOCR()
    
    # Open PDF
    pdf_doc = fitz.open(str(input_path))
    total_pages = len(pdf_doc)
    
    # Create Word document
    doc = Document()
    
    # Process each page
    for page_num in range(total_pages):
        print(f"Page {page_num + 1}/{total_pages}...", end=" ", flush=True)
        
        page = pdf_doc[page_num]
        page_w = page.rect.width
        page_h = page.rect.height
        
        # Render and OCR
        pix = page.get_pixmap(dpi=dpi)
        img_w, img_h = pix.width, pix.height
        
        temp_img = f"temp_page_{page_num}.png"
        pix.save(temp_img)
        
        result, _ = ocr(temp_img)
        
        # Extract elements
        elements = extract_page_elements(result, page_w, page_h, img_w, img_h)
        
        if not elements:
            print("(no text)")
            try:
                os.remove(temp_img)
            except:
                pass
            continue
        
        # Cluster into lines
        lines = cluster_elements_into_lines(elements)
        
        # Calculate average height for heading detection
        avg_height = sum(e['height'] for e in elements) / len(elements)
        
        # Detect table rows
        table_info = detect_table_rows(lines)
        
        print(f"({len(elements)} elements, {len(lines)} lines)")
        
        # Add page separator if not first page
        if page_num > 0:
            doc.add_page_break()
        
        # Process lines and add to document
        i = 0
        while i < len(lines):
            line = lines[i]
            is_table, num_cols = table_info[i][1], table_info[i][2]
            
            # Check if this is start of a table (consecutive table rows)
            if is_table and i + 1 < len(lines) and table_info[i + 1][1]:
                # Collect consecutive table rows
                table_rows = []
                start_i = i
                while i < len(lines) and table_info[i][1]:
                    table_rows.append(lines[i])
                    i += 1
                
                # Determine number of columns (max across all rows)
                max_cols = max(len(row) for row in table_rows)
                
                if max_cols >= 2 and len(table_rows) >= 2:
                    # Create table
                    table = doc.add_table(rows=len(table_rows), cols=max_cols)
                    table.style = 'Table Grid'
                    
                    for row_idx, row in enumerate(table_rows):
                        for col_idx, elem in enumerate(row):
                            if col_idx < max_cols:
                                cell = table.cell(row_idx, col_idx)
                                cell.text = elem['text']
                    
                    doc.add_paragraph()  # Space after table
                    continue
            
            # Regular paragraph/line
            if len(line) == 1:
                elem = line[0]
                text = elem['text']
                
                # Check if heading
                if is_heading(text, elem['height'], avg_height):
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.bold = True
                    run.font.size = Pt(14)
                else:
                    doc.add_paragraph(text)
            else:
                # Multiple elements on same line - join with spaces
                line_text = "  ".join(elem['text'] for elem in line)
                
                # Check for key:value pattern
                if len(line) == 2 and ':' in line[0]['text']:
                    para = doc.add_paragraph()
                    run = para.add_run(line[0]['text'] + " ")
                    run.bold = True
                    para.add_run(line[1]['text'])
                else:
                    doc.add_paragraph(line_text)
            
            i += 1
        
        # Cleanup temp file
        try:
            os.remove(temp_img)
        except:
            pass
    
    pdf_doc.close()
    
    # Save document
    print()
    print("Saving Word document...", end=" ", flush=True)
    doc.save(str(output_path))
    print("done")
    
    print()
    print("=" * 60)
    print(f"Result: {output_path.name}")
    print(f"Pages:  {total_pages}")
    print("=" * 60)
    
    return str(output_path)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        pdf_in = r"C:\Users\harsh\OneDrive\Desktop\Nik\GIG\tes.pdf"
    else:
        pdf_in = sys.argv[1]
    
    docx_out = None
    if len(sys.argv) >= 3:
        docx_out = sys.argv[2]
    
    output = pdf_to_word(pdf_in, docx_out)
    print(f"\nWord document created: {output}")
